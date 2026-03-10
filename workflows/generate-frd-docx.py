"""
FRD Word Document Generator -- converts FRD markdown to styled .docx

Usage:
    python generate-frd-docx.py --input workflows/frd/REQ-001-FRD.md --output workflows/frd/REQ-001-FRD.docx
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_markdown(md_text):
    """Parse FRD markdown into structured sections."""
    sections = []
    current_heading = None
    current_level = 0
    current_lines = []

    for line in md_text.split("\n"):
        h_match = re.match(r'^(#{1,3})\s+(.+)', line)
        if h_match:
            if current_heading is not None:
                sections.append({
                    "heading": current_heading,
                    "level": current_level,
                    "content": "\n".join(current_lines).strip()
                })
            current_heading = h_match.group(2).strip()
            current_level = len(h_match.group(1))
            current_lines = []
        else:
            current_lines.append(line)

    if current_heading is not None:
        sections.append({
            "heading": current_heading,
            "level": current_level,
            "content": "\n".join(current_lines).strip()
        })

    return sections


def add_table_from_md(doc, table_text):
    """Parse markdown table and add to document."""
    rows = []
    for line in table_text.strip().split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r'^\|[\s\-|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)

    if len(rows) < 1:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                cell.text = clean
                for para in cell.paragraphs:
                    para.style.font.size = Pt(10)
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    return table


def add_content(doc, content):
    """Add content block to document, handling tables, lists, bold text."""
    if not content:
        return

    lines = content.split("\n")
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            add_table_from_md(doc, "\n".join(table_lines))
            table_lines = []
            in_table = False

        if stripped == "" or stripped == "---":
            continue

        bullet_match = re.match(r'^[-*]\s+(.+)', stripped)
        if bullet_match:
            para = doc.add_paragraph(style="List Bullet")
            add_formatted_text(para, bullet_match.group(1))
            continue

        num_match = re.match(r'^\d+\.\s+(.+)', stripped)
        if num_match:
            para = doc.add_paragraph(style="List Number")
            add_formatted_text(para, num_match.group(1))
            continue

        para = doc.add_paragraph()
        add_formatted_text(para, stripped)

    if in_table and table_lines:
        add_table_from_md(doc, "\n".join(table_lines))


def add_formatted_text(para, text):
    """Add text with bold/italic markdown formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def create_frd_docx(md_path, docx_path):
    """Convert FRD markdown to styled Word document."""
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    sections = parse_markdown(md_text)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_text = "Functional Requirement Document"
    if sections and sections[0]["level"] == 1:
        title_text = sections[0]["heading"]

    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    start_idx = 1 if (sections and sections[0]["level"] == 1) else 0
    for sec in sections[start_idx:]:
        level = min(sec["level"], 3)
        heading = doc.add_heading(sec["heading"], level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        add_content(doc, sec["content"])

    doc.save(docx_path)
    print(f"Generated: {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="FRD Word Document Generator")
    parser.add_argument("--input", required=True, help="Path to FRD markdown file")
    parser.add_argument("--output", required=True, help="Path for output .docx file")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: {args.input} not found", file=sys.stderr)
        sys.exit(1)

    create_frd_docx(args.input, args.output)


if __name__ == "__main__":
    main()
